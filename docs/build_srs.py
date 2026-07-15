from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("UHLMS_Software_Requirements_Specification.docx")
GREEN = "00491E"
GREEN_2 = "02681E"
GOLD = "FFC600"
PALE_GREEN = "EAF4ED"
PALE_GOLD = "FFF7D6"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
BLACK = "111827"
WHITE = "FFFFFF"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_font(run, size=None, bold=None, italic=None, color=BLACK, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = rgb(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table_borders(table, color="D0D5DD", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MID_GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def add_header_footer(section):
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("UHLMS  |  Software Requirements Specification")
    set_font(r, size=9, bold=True, color=GREEN)
    p_pr = hp._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), GOLD)
    borders.append(bottom)
    p_pr.append(borders)
    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def add_paragraph(doc, text="", style=None, after=6, before=0, keep=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_with_next = keep
    if text:
        r = p.add_run(text)
        set_font(r, size=10.5, color=BLACK)
    return p


def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(f"{label}: ")
    set_font(r, size=10, bold=True, color=GREEN)
    r = p.add_run(value)
    set_font(r, size=10, color=BLACK)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r, size=10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r, size=10.5)
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        shade_cell(hdr.cells[idx], GREEN)
        p = hdr.cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_font(r, size=9.2, bold=True, color=WHITE)
    for row_index, values in enumerate(rows):
        row = table.add_row()
        if row_index % 2:
            for cell in row.cells:
                shade_cell(cell, "F8FAFC")
        for idx, value in enumerate(values):
            p = row.cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            if idx == 0 and len(headers) <= 3:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_font(r, size=8.8, color=BLACK)
    set_table_geometry(table, widths)
    table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, title, text, fill=PALE_GREEN):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), GOLD)
    borders.append(left)
    p_pr.append(borders)
    r = p.add_run(title)
    set_font(r, size=10.5, bold=True, color=GREEN)
    r = p.add_run("\n" + text)
    set_font(r, size=10, color=BLACK)


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, before, after, color in (
        ("Title", 28, 0, 8, GREEN),
        ("Subtitle", 13, 0, 12, MID_GRAY),
        ("Heading 1", 16, 18, 10, GREEN),
        ("Heading 2", 13, 14, 7, GREEN_2),
        ("Heading 3", 11.5, 10, 5, GREEN),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


USE_CASES = [
    ("UC-01", "Access the Public Guest Site", "Public Information and Discovery", "Visitor, Guest", "Essential", "Always", "Open the UHLMS homepage and access public navigation, notices, property information, contact details, and calls to action.", ["The actor opens the localhost or Cloudflare-hosted site.", "The system renders the current guest-site branding, announcements, navigation, room summary, stay guide, and contact details.", "The actor selects a public destination or begins a stay request."], "The public experience must remain usable in both supported runtime targets."),
    ("UC-02", "Check Quick Availability", "Public Information and Discovery", "Visitor, Guest", "Essential", "Often", "Check room availability using check-in, check-out, and guest-count criteria from the homepage.", ["The actor enters valid dates and guest count.", "The system validates the date range.", "The system opens the room catalog with the criteria applied and displays availability summaries."], "Availability is advisory until staff review and room holding."),
    ("UC-03", "Search and Filter Room Catalog", "Public Information and Discovery", "Visitor, Guest", "Essential", "Often", "Browse active room types and narrow results by dates, capacity, amenities, sharing type, pricing type, budget, availability visibility, and sort order.", ["The actor opens the room catalog.", "The actor applies one or more filters.", "The system validates criteria, calculates date-aware capacity, and returns matching room types.", "The actor may include or hide currently unavailable results."], "Shared rooms are evaluated by remaining bed capacity; private rooms by available rooms and capacity."),
    ("UC-04", "View Room Details", "Public Information and Discovery", "Visitor, Guest", "Essential", "Often", "View a room type's description, rate, capacity, amenities, images, availability context, and linked virtual-tour scene.", ["The actor selects a room type.", "The system displays active room details and preserves availability parameters.", "The actor may open the relevant tour scene or continue to request a stay."], "Request links must use request-aware relative routing."),
    ("UC-05", "Explore the Virtual Tour", "Virtual Tour", "Visitor, Guest", "Expected", "Often", "Navigate active 360-degree scenes, panoramas, hotspots, room information cards, and guided transitions.", ["The actor opens the virtual tour.", "The system loads the requested active scene or the default scene.", "The actor pans, zooms, enters supported VR/stereo modes, and follows hotspots.", "The system loads linked scene or media content."], "Viewer behavior depends on browser/device media and motion capabilities."),
    ("UC-06", "Check Tour Availability", "Virtual Tour", "Visitor, Guest", "Essential", "Often", "Retrieve current or date-aware room and room-type availability while exploring the tour.", ["The actor opens room information in a tour scene.", "The client requests availability from the tour API.", "The system returns room, room-type, capacity, and request guidance.", "The viewer updates the card and available request actions."], "The API must not expose private reservation details."),
    ("UC-07", "Register a Guest Account", "Guest Identity and Access", "Visitor", "Essential", "Sometimes", "Create a guest account with identity, contact information, and password.", ["The visitor opens registration and supplies required information.", "The system validates uniqueness, mobile format, password rules, and anti-spam controls.", "The system creates and signs in the account.", "The system sends an email-verification link."], "Unverified accounts have restricted feedback eligibility and receive verification guidance."),
    ("UC-08", "Verify Guest Email", "Guest Identity and Access", "Guest, Email Service", "Essential", "Once", "Verify control of the email address associated with a guest account.", ["The guest opens the verification link.", "The system validates the account reference and link.", "The system records the verification timestamp and confirms the result."], "A signed or otherwise integrity-protected link is required."),
    ("UC-09", "Authenticate and End Guest Session", "Guest Identity and Access", "Guest", "Essential", "Often", "Sign in to and sign out from the guest account area.", ["The guest enters credentials.", "The system validates that the account exists, is enabled, and the password is correct.", "The system regenerates the authenticated session and records login activity.", "On logout, the system invalidates the session."], "Disabled accounts must not gain authenticated access."),
    ("UC-10", "Reset Guest Password", "Guest Identity and Access", "Guest, Email Service", "Essential", "Rarely", "Request and complete a guest password reset through email.", ["The guest submits the account email.", "The system issues a time-limited reset link without disclosing account existence unnecessarily.", "The guest supplies and confirms a new password.", "The system validates the token and updates the password."], "Reset tokens must expire and be single-purpose."),
    ("UC-11", "Manage Guest Profile", "Guest Identity and Access", "Guest", "Expected", "Sometimes", "Review and update permitted guest profile and contact information.", ["The authenticated guest opens the profile page.", "The system loads current data.", "The guest edits permitted fields.", "The system validates and saves the changes."], "Identity and contact changes must not expose another guest's records."),
    ("UC-12", "Submit a Stay Request", "Stay Request and Tracking", "Visitor, Guest", "Essential", "Sometimes", "Submit a pending stay request for one or multiple room types, dates, guest count, purpose, and special requests.", ["The actor opens the request form, optionally prefilled from room search.", "The actor provides personal and stay information and selects one or more room types.", "The system validates dates, age, contact format, capacity, duplication, and low-availability acknowledgement.", "The system creates the reservation and requested-room lines with a unique reference.", "The system sends confirmation/tracking guidance and redirects to tracking."], "Submission is a request, not an immediate booking; normal review is communicated as 1-2 business days."),
    ("UC-13", "Request a Room from the Tour", "Stay Request and Tracking", "Visitor, Guest", "Essential", "Sometimes", "Submit a room-type request directly from the virtual-tour interface.", ["The actor opens a linked room scene.", "The actor enters stay and guest information and acknowledges limited availability when required.", "The tour API validates and creates a pending reservation request.", "The system returns a reference and next-step guidance."], "Rate limiting and honeypot protection apply."),
    ("UC-14", "Track a Reservation", "Stay Request and Tracking", "Guest, Visitor", "Essential", "Often", "View reservation status and next-step guidance using reference plus email or a signed secure link.", ["The actor supplies the reference and matching email, or opens a valid signed link.", "The system validates lookup authority and retention rules.", "The system displays status, essential dates, requested room summary, payment eligibility, and guidance."], "Cancelled/declined and completed tracking records expire according to configured retention behavior."),
    ("UC-15", "Use Guest Dashboard", "Guest Self-Service", "Guest", "Expected", "Often", "View categorized reservation statistics, recent stays, pending actions, verification status, and account navigation.", ["The guest opens the dashboard.", "The system loads only reservations linked to the account.", "The system groups reservations into pending, alternative-offer, active, completed, and upcoming views.", "The guest opens an available action."], "Dashboard counts and labels must remain aligned with reservation status definitions."),
    ("UC-16", "Claim an Existing Reservation", "Guest Self-Service", "Guest", "Expected", "Rarely", "Link a previously submitted reservation to the authenticated guest account.", ["The guest enters the reservation reference and matching identifying information.", "The system validates ownership and claim eligibility.", "The system links the reservation to the account.", "The reservation appears in the guest dashboard."], "A reservation already owned by another account cannot be claimed."),
    ("UC-17", "View Guest Reservation Details", "Guest Self-Service", "Guest", "Essential", "Often", "View a private, account-scoped reservation record including status, requested rooms, payments, assignments when permitted, and feedback state.", ["The guest selects a linked reservation.", "The system verifies ownership.", "The system loads related room requests, payments, holds/assignments, and feedback.", "The system displays applicable next actions."], "Sensitive internal notes and unrelated guest details must not be exposed."),
    ("UC-18", "Respond to Alternative Room Offer", "Guest Self-Service", "Guest", "Essential", "Sometimes", "Accept or decline a time-limited alternative room offer proposed by staff.", ["The guest opens the offer link.", "The system validates offer status and expiry.", "The guest reviews offered room type, held capacity, dates, and message.", "On acceptance, the system confirms the alternative allocation and updates the reservation; on decline, it releases holds and records the response."], "Expired offers cannot be accepted and are automatically reconciled."),
    ("UC-19", "Make an Online Payment", "Payments", "Guest, PayMongo", "Expected", "Sometimes", "Pay an eligible deposit or balance using enabled online payment methods.", ["The guest opens a valid, unexpired payment link for an eligible reservation.", "The system calculates the payable amount and displays configured methods.", "The guest chooses a method and is redirected to or completes the gateway flow.", "The gateway returns a success, pending, or failed outcome."], "Online payment is feature-controlled and requires an active advance room hold."),
    ("UC-20", "Process Payment Result", "Payments", "System, PayMongo", "Essential", "Sometimes", "Verify gateway callbacks/webhooks, record payment data, update totals/status, and confirm eligible reservations.", ["PayMongo sends a signed event or the client returns with a gateway result.", "The system verifies authenticity and idempotency.", "The system records gateway identifiers and payment status.", "Successful payment updates reservation financials and may mark the reservation confirmed.", "The system displays or emails the resulting status."], "Webhook processing must not trust unsigned payloads and must tolerate retries."),
    ("UC-21", "Submit a Support Inquiry", "Support", "Visitor, Guest", "Expected", "Sometimes", "Send a categorized question, request, complaint, or feedback item to homestay staff.", ["The actor opens support.", "An authenticated guest uses the support inbox; a visitor uses the public support form.", "The actor supplies required contact, category, subject, and message data.", "The system validates, rate-limits, stores, and acknowledges the inquiry."], "Authenticated inquiries are linked to the guest account for threaded replies."),
    ("UC-22", "Continue a Support Conversation", "Support", "Guest, Staff", "Expected", "Sometimes", "Exchange replies on an account-linked support inquiry.", ["The guest or authorized staff opens an inquiry.", "The system loads the conversation in chronological order.", "The actor submits a reply.", "The system validates ownership/permission, stores the message, and refreshes the thread."], "Only authorized participants may read or reply to a thread."),
    ("UC-23", "Submit Stay Feedback", "Feedback and Testimonials", "Verified Guest", "Expected", "Sometimes", "Rate a completed stay and optionally provide category ratings, comments, return intent, public-comment consent, and separate room-type display consent.", ["A verified guest opens feedback for an owned checked-out reservation.", "The guest enters ratings and optional comments.", "The guest independently chooses whether the comment may be considered for public display and whether room type may also be shown.", "The system validates and stores one internal feedback record for the reservation."], "Feedback is internal by default and cannot be submitted twice for the same reservation."),
    ("UC-24", "Display Public Testimonials", "Feedback and Testimonials", "Visitor, System", "Expected", "Often", "Display anonymous testimonials that satisfy consent and staff-review requirements.", ["The homepage requests eligible testimonials.", "The system selects only reviewed, explicitly public, consented feedback with written comments.", "The system displays rating, comment, and 'Verified guest' attribution.", "The room type appears only when the guest separately consented to its display."], "Names, email, reservation reference, room number, and stay dates remain private."),
    ("UC-25", "Authenticate Staff and Manage Profile", "Staff Access and Dashboard", "Staff, Administrator, Super Administrator", "Essential", "Always", "Access the Filament administration panel, reset password, and update the staff profile.", ["The user signs in with a staff account.", "The system verifies credentials and an allowed role.", "The system opens the panel with permission-filtered navigation.", "The user may update permitted profile information or end the session."], "Panel access is limited to staff, administrator, and super-administrator roles."),
    ("UC-26", "Monitor Operational Dashboard", "Staff Access and Dashboard", "Staff, Administrator", "Essential", "Often", "View reservation counts and operational indicators such as pending review, approved/confirmed stays, current occupancy, near-due checkouts, and overdue stays.", ["The user opens the admin dashboard.", "The system calculates authorized operational metrics.", "The user selects a metric to open the corresponding filtered records."], "Metrics must use the same status definitions as reservation records and reports."),
    ("UC-27", "Create and Edit Reservation", "Reservation Operations", "Authorized Staff", "Essential", "Often", "Create walk-in/staff-entered reservations and edit permitted reservation, guest, room-request, discount, service, and financial details.", ["The user opens create or edit.", "The system presents fields allowed by permission and reservation state.", "The user enters or modifies data.", "The system validates and persists the record with audit timestamps."], "Edits must respect active holds, assignments, and status-transition rules."),
    ("UC-28", "Approve Reservation and Hold Rooms", "Reservation Operations", "Authorized Staff", "Essential", "Often", "Approve a pending request, select appropriate available rooms by requested type, and create advance holds.", ["The user reviews a pending reservation and current availability.", "The user selects rooms by requested room type and optional notes.", "The system revalidates availability and creates time-limited advance holds.", "The system updates status to approved, creates/refreshes payment access when enabled, logs the event, and notifies the guest."], "Approval must not bypass room-conflict checks."),
    ("UC-29", "Propose Alternative Accommodation", "Reservation Operations", "Authorized Staff", "Expected", "Sometimes", "Offer a different room type when the requested accommodation is unavailable.", ["The user selects an unavailable pending request.", "The system lists eligible alternative rooms and capacity.", "The user selects rooms, hold duration, and optional message.", "The system creates the offer and holds, updates status, and emails the guest."], "Only one active offer should govern a reservation; superseded or expired holds are released."),
    ("UC-30", "Decline or Cancel Reservation", "Reservation Operations", "Authorized Staff", "Essential", "Sometimes", "Decline a request or cancel an existing reservation with a recorded reason.", ["The user selects decline or cancel and supplies the required reason when applicable.", "The system validates the transition.", "The system updates status, releases holds, invalidates payment access as required, logs the action, and notifies the guest."], "The original record remains available for audit and reporting unless separately force-deleted."),
    ("UC-31", "Check In Guests", "Stay Operations", "Authorized Staff", "Essential", "Often", "Finalize guest identity, room assignments, companions, ID details, discounts, add-ons, payment, and check-in snapshot.", ["The user opens check-in for an eligible reservation or onsite flow.", "The system loads held/available rooms and expected guest data.", "The user records primary and companion guests, assignments, ID, dates, services, discount evidence, payment, and remarks.", "The system validates capacity and financial data, creates assignments/charges/payments/snapshot/logs, clears holds, and marks the reservation checked in."], "All assigned guests must fit room capacity and no room conflict may be introduced."),
    ("UC-32", "Check Out Guests", "Stay Operations", "Authorized Staff", "Essential", "Often", "Complete a stay, release rooms, and record checkout timing and remarks.", ["The user selects a checked-in reservation.", "The user confirms checkout date/time and optional remarks.", "The system updates assignment and room states, records stay logs, updates the reservation to checked out, and recalculates availability."], "Forced bulk checkout is restricted and audited."),
    ("UC-33", "Manage Reservation Payments and Charges", "Payments", "Authorized Staff", "Essential", "Often", "Record onsite/manual payments, official receipts, charges, add-ons, payment modes, and balances.", ["The user opens a reservation's payment or check-in financial section.", "The system displays current charges, payments, and balance.", "The user records valid payment and receipt information.", "The system persists financial records and recalculates totals and payment status."], "Financial changes require reservation-edit permission and must retain gateway references when applicable."),
    ("UC-34", "Manage Payment Links", "Payments", "Authorized Staff", "Expected", "Sometimes", "Refresh, email, copy, and inspect a guest payment link and QR code for an eligible reservation.", ["The user selects payment-link management.", "The system validates online-payment configuration and reservation eligibility.", "The system issues a new token/expiry and optionally emails it.", "The user may copy the request-aware URL or display the encrypted QR code."], "Refreshing invalidates the usefulness of expired or superseded access."),
    ("UC-35", "Manage Room Holds", "Availability and Inventory", "Authorized Staff", "Essential", "Often", "Review active holds by room/date/reservation and release permitted holds individually or in bulk.", ["The user opens room holds and filters by activity or date.", "The system displays hold type, capacity, reservation, and expiry.", "An authorized user releases selected advance holds.", "The system recalculates availability."], "Release permission is separate from view permission."),
    ("UC-36", "Manage Rooms", "Availability and Inventory", "Administrator", "Essential", "Sometimes", "Create, view, update, activate/deactivate, and delete physical room records with number, type, floor, capacity, and operational status.", ["The user opens Rooms.", "The system enforces the applicable CRUD permission.", "The user enters or edits room information.", "The system validates uniqueness and relationships and saves the record."], "Destructive changes must not leave active assignments or holds inconsistent."),
    ("UC-37", "Manage Room Types", "Availability and Inventory", "Administrator", "Essential", "Sometimes", "Manage room categories, rates, pricing model, sharing model, capacity, description, images, and amenity associations.", ["The user opens Room Types.", "The user creates or edits a type and related attributes.", "The system validates pricing/capacity and saves relationships.", "Active types become eligible for public catalog and availability use."], "Public sharing and per-person pricing must be distinguished from private flat-rate types."),
    ("UC-38", "Manage Floors", "Availability and Inventory", "Administrator", "Expected", "Rarely", "Maintain floor names, levels, descriptions, and active state used by rooms.", ["The user opens Floors.", "The user creates or edits a floor.", "The system validates unique level/name rules and persists the record."], "A floor referenced by rooms cannot be removed without preserving integrity."),
    ("UC-39", "Manage Amenities", "Availability and Inventory", "Administrator", "Expected", "Sometimes", "Maintain active amenities and associate them with room types.", ["The user opens Amenities.", "The user creates, edits, or deletes an amenity subject to permission.", "The system saves associations and updates guest-facing filters/details."], "Only active amenities appear publicly."),
    ("UC-40", "Manage Add-On Services", "Availability and Inventory", "Administrator", "Expected", "Sometimes", "Maintain add-on service names, prices, units, descriptions, display order, and active state.", ["The user opens Add-Ons.", "The user creates or edits a service.", "The system validates price and saves it.", "Active services become available in stay guidance and staff check-in pricing."], "Historical reservation charges must retain their recorded values."),
    ("UC-41", "Manage Guest Accounts", "Guest Administration", "Authorized Staff", "Expected", "Sometimes", "View guest profiles and reservation history; edit, resend verification, disable, or enable accounts according to permission.", ["The user opens Guest Accounts.", "The system shows account status, verification, reservation counts, and latest stay.", "The user performs an authorized action.", "The system validates permission, updates the account, and records the result."], "View, edit, and disable/enable are distinct permissions."),
    ("UC-42", "Moderate Guest Feedback", "Feedback and Testimonials", "Authorized Staff", "Expected", "Sometimes", "Review feedback, record internal notes, mark reviewed, and manage eligible public-testimonial visibility.", ["The user opens feedback and filters by rating, status, date, or room type.", "The user reviews ratings, comment, return intent, and guest consent flags.", "The user marks the record reviewed and may approve it for public display only when consent and written-comment conditions are satisfied.", "The user may remove a testimonial from public display."], "The model must force ineligible feedback to internal visibility."),
    ("UC-43", "Triage and Reply to Support Inquiries", "Support", "Authorized Staff", "Essential", "Often", "Review, categorize, assign status, add internal notes, and reply to guest-linked support inquiries.", ["The user opens Support Inquiries or the support inbox.", "The system displays contact, category, message, account link, status, and conversation.", "The user updates triage data or submits a reply.", "The system records handler/reply metadata and exposes the reply to the authorized guest."], "Internal notes are never exposed in the guest conversation."),
    ("UC-44", "Manage Virtual-Tour Scenes", "Virtual Tour Administration", "Authorized Administrator", "Expected", "Sometimes", "Create, edit, preview, activate/deactivate, and delete panorama scenes with thumbnails, default view, room-type links, and ordering.", ["The user opens Virtual Tours.", "The user uploads a panorama and enters scene metadata.", "The system validates media and saves the scene.", "The user previews or activates the scene for public use."], "View and manage permissions are separate; media URLs must work locally and through Cloudflare."),
    ("UC-45", "Manage Tour Hotspots", "Virtual Tour Administration", "Authorized Administrator", "Expected", "Sometimes", "Create and position scene-navigation, media, informational, room, previous-scene, and other supported hotspots.", ["The user opens hotspot management for a scene.", "The editor loads the panorama and existing hotspots.", "The user adds/edits action type, target, label, position, size, and media.", "The system validates and persists hotspot configuration."], "Broken targets must be prevented or surfaced clearly."),
    ("UC-46", "Configure Guest Site", "System Administration", "Administrator", "Expected", "Sometimes", "Configure branding, theme, notices, hero, sections, policies, FAQs, request guidance, navigation, footer, accessibility modes, and media.", ["The user opens Guest Site Settings.", "The system loads defaults plus stored overrides.", "The user edits permitted settings and media.", "The system validates and saves values.", "Public pages reflect the changes."], "Settings and media URLs must remain compatible with localhost and Cloudflare."),
    ("UC-47", "Configure Online Payments", "System Administration", "Administrator", "Expected", "Rarely", "Enable/disable online payments, set default deposit percentage, inspect PayMongo key readiness, and obtain the webhook URL.", ["The user opens Online Payment Settings.", "The system displays feature and integration status without revealing secrets.", "The user changes permitted settings.", "The system validates and saves configuration."], "Gateway secrets are environment-managed and must not be displayed or stored in plain guest settings."),
    ("UC-48", "Configure Reservation Discounts", "System Administration", "Administrator", "Expected", "Rarely", "Set PWD, senior citizen, and student discount rates used during reservation/check-in calculations.", ["The user opens Discount Settings.", "The user supplies valid percentage values.", "The system validates ranges and saves them.", "Subsequent eligible pricing uses the configured rates."], "Declared discounts remain subject to staff verification at check-in."),
    ("UC-49", "Configure Report Signatories", "System Administration", "Administrator", "Expected", "Rarely", "Maintain prepared-by and approved-by names/titles used in formal report outputs.", ["The user opens Report Signatories.", "The system loads current names and positions.", "The user edits and saves them.", "Generated reports use the updated signatory data."], "Blank signatories should render gracefully."),
    ("UC-50", "Manage Staff Users and Permissions", "System Administration", "Administrator, Super Administrator", "Essential", "Sometimes", "Create and manage staff accounts, roles, role defaults, and per-user custom permissions.", ["The authorized user opens Users.", "The system applies role-based visibility and protection rules.", "The user creates or edits account credentials, role, or custom permission toggles.", "The system validates and saves the account.", "Navigation and actions reflect the effective permission set."], "Only a super administrator may manage protected administrator-level attributes and reset custom permissions to role defaults."),
    ("UC-51", "Review Permission Reference", "System Administration", "Administrator, Super Administrator", "Expected", "Sometimes", "View the centralized matrix of permissions grouped by functional module.", ["The user opens the permission reference from Users.", "The system displays permission keys, labels, and role-default context.", "The user uses the matrix when configuring access."], "The reference must remain aligned with the permission definitions used by authorization checks."),
    ("UC-52", "Back Up and Restore the System", "System Administration", "Administrator", "Essential", "Rarely", "Create/download backups, upload a supported backup, and restore application data through the guarded administration workflow.", ["The user opens Backup and Restore.", "The user creates or selects a backup action.", "The system validates authorization, file type, and operation prerequisites.", "The system performs the backup/restore and reports success or failure."], "Restore is a high-impact action and must be restricted, validated, and logged."),
    ("UC-53", "Review Force-Deletion Logs", "System Administration", "Super Administrator", "Essential", "Rarely", "Audit privileged reservation force-deletion and force-checkout activity, including actor and reason.", ["The user opens Force Deletion Logs.", "The system verifies privileged access.", "The system displays immutable deletion/force-action evidence and related identifiers."], "Deletion evidence must survive the deletion of the operational record where designed."),
    ("UC-54", "View Room Utilization Calendar", "Reporting and Analytics", "Authorized Staff", "Expected", "Often", "View date-based room occupancy, holds, assignments, and reservation status in a calendar-oriented operational display.", ["The user opens the utilization calendar and selects a date/range.", "The system compiles room, assignment, hold, and status data.", "The calendar displays availability and occupancy using the shared status palette.", "The user opens relevant records when permitted."], "Legend colors must remain consistent with reservation status presentation."),
    ("UC-55", "Generate Monthly OR Report", "Reporting and Analytics", "Authorized Staff", "Expected", "Monthly", "Generate the monthly official-receipt/payment-oriented management report and export it when separately permitted.", ["The user selects the reporting period.", "The system aggregates qualifying payment and reservation data.", "The system displays totals and detailed rows.", "An authorized exporter generates the supported downloadable output with signatories."], "View and export are separate permissions."),
    ("UC-56", "Generate Reservation Summary", "Reporting and Analytics", "Authorized Staff", "Expected", "Often", "Summarize reservations for a period by status and room type.", ["The user selects a reporting range.", "The system aggregates reservations by status and preferred/requested room type.", "The system displays totals and breakdowns."], "Counts must use canonical reservation statuses."),
    ("UC-57", "Generate Reservation List", "Reporting and Analytics", "Authorized Staff", "Expected", "Often", "Produce a detailed period-based reservation list with guest, stay, room, status, and related operational fields.", ["The user selects date criteria and permitted filters.", "The system retrieves matching reservations.", "The system displays a sortable/report-ready list."], "The report must respect authorization and privacy constraints."),
    ("UC-58", "Generate Occupancy Report", "Reporting and Analytics", "Authorized Staff", "Expected", "Often", "Analyze occupied capacity and occupancy rates over a selected period.", ["The user selects a period.", "The system calculates room nights/capacity and occupied values from assignments and stays.", "The system displays occupancy totals and rates."], "The report must distinguish rooms, beds, and capacity as appropriate to sharing type."),
    ("UC-59", "Generate Room Utilization Report", "Reporting and Analytics", "Authorized Staff", "Expected", "Often", "Analyze room and room-type utilization using assignments and stay duration.", ["The user selects a period.", "The system groups assignment/stay activity by room and room type.", "The system displays utilization counts, duration, and related measures."], "Cancelled/declined requests must not be counted as occupied stays."),
    ("UC-60", "Generate Gender Statistics", "Reporting and Analytics", "Authorized Staff", "Expected", "Sometimes", "Summarize recorded guest gender data for an authorized reporting period.", ["The user selects a period.", "The system aggregates available primary/companion guest gender records.", "The system displays counts and proportions."], "Results should include unknown/unspecified data rather than silently discarding it."),
    ("UC-61", "Generate Feedback Analytics", "Reporting and Analytics", "Authorized Staff", "Expected", "Sometimes", "Analyze feedback volume, average rating, distribution, category ratings, trends, room-type results, low ratings, and review status.", ["The user selects a reporting period.", "The system loads eligible feedback and related room types.", "The system calculates summary, distribution, trend, category, repeat-intent, low-rating, and unreviewed measures.", "The system displays the analytics report."], "Public visibility is not required for feedback to contribute to internal analytics."),
    ("UC-62", "Generate Stay Logs Report", "Reporting and Analytics", "Authorized Staff", "Essential", "Sometimes", "Review check-in, checkout, assignment, processor, timing, and remarks history for completed/active stays.", ["The user selects a period.", "The system retrieves authorized stay and assignment logs.", "The system displays traceable operational history."], "Stay-log viewing is permission-controlled and preserves audit context."),
    ("UC-63", "Send Reservation and Account Email", "Automated Services", "System, Email Service", "Essential", "Often", "Send verification, password reset, reservation status, tracking, alternative offer, payment-link, and related transactional emails.", ["A qualifying account or reservation event occurs.", "The system builds the appropriate request-aware or signed link.", "The system queues/sends the email.", "Delivery failure is logged and the core transaction remains consistent."], "Links must resolve correctly for both local testing and the Cloudflare hostname."),
    ("UC-64", "Run Scheduled Expiration and Cleanup", "Automated Services", "System Scheduler", "Essential", "Hourly", "Expire alternative offers and related holds, and reconcile reservation state without manual intervention.", ["The scheduler runs the expiration command.", "The system identifies open offers past expiry.", "The system marks them expired, releases held inventory, and restores the appropriate reservation state.", "The system logs the outcome."], "The operation must be idempotent and safe to retry."),
]


FUNCTIONAL_GROUPS = [
    ("Public Information and Discovery", "UC-01 to UC-04"),
    ("Virtual Tour", "UC-05 to UC-06"),
    ("Guest Identity and Access", "UC-07 to UC-11"),
    ("Stay Request and Tracking", "UC-12 to UC-14"),
    ("Guest Self-Service", "UC-15 to UC-18"),
    ("Payments", "UC-19 to UC-20, UC-33 to UC-34"),
    ("Support", "UC-21 to UC-22, UC-43"),
    ("Feedback and Testimonials", "UC-23 to UC-24, UC-42"),
    ("Staff Access and Dashboard", "UC-25 to UC-26"),
    ("Reservation and Stay Operations", "UC-27 to UC-32"),
    ("Availability and Inventory", "UC-35 to UC-40"),
    ("Guest Administration", "UC-41"),
    ("Virtual Tour Administration", "UC-44 to UC-45"),
    ("System Administration", "UC-46 to UC-53"),
    ("Reporting and Analytics", "UC-54 to UC-62"),
    ("Automated Services", "UC-63 to UC-64"),
]


STAKEHOLDERS = [
    ("Visitor / Prospective Guest", "UC-01 to UC-07, UC-12 to UC-14, UC-21, UC-24"),
    ("Registered Guest", "UC-08 to UC-24"),
    ("Staff", "UC-25 to UC-35, UC-41 to UC-43, UC-54 to UC-62"),
    ("Administrator", "All staff use cases plus UC-36 to UC-52"),
    ("Super Administrator", "All administrative use cases, protected user/permission operations, and UC-53"),
    ("PayMongo", "UC-19 to UC-20"),
    ("Email Service", "UC-08, UC-10, UC-18 to UC-20, UC-22, UC-28 to UC-30, UC-34, UC-41, UC-63"),
    ("System Scheduler", "UC-64"),
]


def build_document():
    doc = Document()
    setup_styles(doc)
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)
        add_header_footer(section)

    # Cover - editorial_cover pattern with CMU identity override.
    doc.add_paragraph().paragraph_format.space_after = Pt(54)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CENTRAL MINDANAO UNIVERSITY")
    set_font(r, size=11, bold=True, color=GREEN)
    p.paragraph_format.space_after = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Software Requirements Specification")
    set_font(r, size=28, bold=True, color=GREEN)
    p.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("University Homestay Lodging Management System (UHLMS)")
    set_font(r, size=16, bold=True, color=GREEN_2)
    p.paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Complete functional and non-functional requirements baseline")
    set_font(r, size=11, italic=True, color=MID_GRAY)
    p.paragraph_format.space_after = Pt(54)
    add_table(doc, ["Document control", "Value"], [
        ("System", "University Homestay Lodging Management System"),
        ("Document type", "Software Requirements Specification"),
        ("Version", "1.0"),
        ("Status", "Draft for review and validation"),
        ("Prepared", date.today().strftime("%B %d, %Y")),
        ("Runtime targets", "Localhost/LAN and Cloudflare-hosted web application"),
    ], [2400, 6960])
    add_callout(doc, "Document purpose", "This SRS consolidates the implemented and expected behavior of UHLMS into a traceable use-case baseline for development, testing, deployment, evaluation, and future maintenance.", PALE_GOLD)
    doc.add_page_break()

    doc.add_heading("Revision History", level=1)
    add_table(doc, ["Version", "Date", "Description", "Author/Reviewer"], [
        ("1.0", date.today().isoformat(), "Initial complete SRS baseline derived from the implemented UHLMS feature set and supplied use-case layout reference.", "Project Team"),
    ], [1100, 1700, 4760, 1800])

    doc.add_heading("Contents", level=1)
    for item in [
        "1. Introduction", "2. Overall Description", "3. External Interface Requirements",
        "4. Use-Case Catalog", "5. Functional Requirements", "6. Data Requirements",
        "7. Business Rules", "8. Non-Functional Requirements", "9. Traceability and Acceptance",
        "Appendix A. Reservation Status Lifecycle", "Appendix B. Glossary",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()

    doc.add_heading("1. Introduction", level=1)
    doc.add_heading("1.1 Purpose", level=2)
    add_paragraph(doc, "This document specifies the software requirements for the University Homestay Lodging Management System (UHLMS). It defines the system boundary, actors, interfaces, data, business rules, functional behavior, quality attributes, and acceptance expectations for the guest-facing website and staff administration panel.")
    doc.add_heading("1.2 Scope", level=2)
    add_paragraph(doc, "UHLMS supports discovery of university homestay accommodations, date-aware availability, virtual tours, stay requests, guest accounts, reservation tracking, alternative room offers, payments, check-in/check-out, inventory, feedback, support, public testimonials, reporting, backups, and role-based administration. It is a responsive Laravel web application supported on localhost/LAN and through the named Cloudflare-hosted environment.")
    doc.add_heading("1.3 Intended Audience", level=2)
    for text in ["Project advisers, evaluators, and institutional stakeholders validating system scope.", "Developers and maintainers implementing or extending UHLMS.", "Testers deriving functional, integration, security, and acceptance tests.", "Homestay staff and administrators validating operational workflows.", "Deployment personnel supporting local and Cloudflare runtime targets."]:
        add_bullet(doc, text)
    doc.add_heading("1.4 Definitions and Conventions", level=2)
    add_paragraph(doc, "The keyword shall denotes a mandatory requirement; should denotes a recommended quality or operational expectation; may denotes an optional capability. Use-case identifiers are stable requirement references. 'Guest' means an authenticated guest account unless the use case explicitly includes an unauthenticated visitor.")
    doc.add_heading("1.5 Reference Basis", level=2)
    add_paragraph(doc, "The structure follows the supplied use-case reference: use cases are grouped by functional area and stakeholder, then specified with summary, priority, use frequency, direct actors, main success scenario, and notes. The content is adapted to the current UHLMS routes, controllers, services, data models, administrative resources, permissions, reports, and recently implemented testimonial/privacy behavior.")

    doc.add_heading("2. Overall Description", level=1)
    doc.add_heading("2.1 Product Perspective", level=2)
    add_paragraph(doc, "UHLMS is a centralized web application consisting of a public/guest experience, a guest-account self-service area, an administrative panel, scheduled/background processes, and external integrations. A single relational data model coordinates rooms, room types, requests, holds, assignments, guests, payments, feedback, support, settings, and audit records.")
    doc.add_heading("2.2 Product Functions", level=2)
    functions = [
        "Publish property, room, rate, amenity, stay-guide, policy, FAQ, and testimonial content.",
        "Calculate real-time and date-aware availability for private rooms and shared-bed inventory.",
        "Present linked 360-degree virtual-tour scenes and interactive hotspots.",
        "Accept single- and multi-room-type stay requests from the form or tour.",
        "Track reservation status by verified lookup, signed link, or guest account.",
        "Manage approval, alternatives, holds, room assignments, check-in, checkout, charges, payments, and logs.",
        "Support guest registration, verification, password reset, profile, reservation claiming, feedback, and support.",
        "Manage room inventory, site content, payment/discount/report settings, users, permissions, backups, and media.",
        "Generate operational and analytical reports with permission-controlled access.",
    ]
    for text in functions:
        add_bullet(doc, text)
    doc.add_heading("2.3 User Classes", level=2)
    add_table(doc, ["User class", "Description", "Typical access"], [
        ("Visitor", "Unauthenticated prospective guest or public viewer.", "Public pages, rooms, tour, request, tracking, public support."),
        ("Registered Guest", "Guest with an account; verification may affect eligibility.", "Dashboard, profile, owned reservations, feedback, support conversations."),
        ("Staff", "Operational user with role defaults or custom permissions.", "Reservations, check-in/out, holds, support, permitted reports and view-only inventory."),
        ("Administrator", "Management user with broad operational/configuration access.", "CRUD inventory, settings, users, reports, feedback moderation, tour management."),
        ("Super Administrator", "Highest-privilege user.", "All permissions plus protected user/role and destructive audit operations."),
        ("External Services", "PayMongo, mail transport, scheduler, and Cloudflare transport.", "Defined APIs, callbacks, email delivery, scheduled execution, HTTPS routing."),
    ], [1700, 3760, 3900])
    doc.add_heading("2.4 Operating Environment", level=2)
    for text in ["Server: PHP/Laravel application with relational database, queue/cache/session services as configured.", "Client: modern desktop and mobile browsers; virtual-tour capabilities vary by browser/device.", "Administrative UI: Filament-based web panel.", "Runtime targets: localhost/LAN and https://app.uhlms.uk through Cloudflare Tunnel.", "External services: SMTP-compatible mail and PayMongo when online payments are enabled."]:
        add_bullet(doc, text)
    doc.add_heading("2.5 Constraints and Assumptions", level=2)
    for text in ["A submitted stay request does not guarantee accommodation until staff review and room allocation.", "Room and bed availability depends on active rooms, assignments, date overlap, and active holds.", "Public testimonial display requires written feedback, guest consent, completed staff review, and explicit public approval.", "Online payment requires enabled settings, valid environment credentials, an eligible reservation, and active advance hold.", "System-generated links must be request-aware or relative where possible and valid in both supported runtime targets.", "Email delivery and payment gateway availability are external dependencies."]:
        add_bullet(doc, text)

    doc.add_heading("3. External Interface Requirements", level=1)
    doc.add_heading("3.1 User Interfaces", level=2)
    add_paragraph(doc, "The public interface shall provide responsive navigation, accessible labels, validation messages, date and numeric controls, status guidance, and mobile-friendly cards/forms. The administrative interface shall expose only authorized navigation, pages, fields, row actions, bulk actions, and reports. Destructive or public-visibility actions shall require deliberate confirmation where implemented.")
    doc.add_heading("3.2 Software Interfaces", level=2)
    add_table(doc, ["Interface", "Direction", "Purpose", "Primary controls"], [
        ("PayMongo API", "Outbound/inbound", "Checkout/payment creation, retrieval, and webhooks.", "HTTPS, secrets in environment, signature verification, idempotent processing."),
        ("Email service", "Outbound", "Verification, password reset, tracking, offers, status and payment messages.", "Time-limited/signed links, failure logging, no secret disclosure."),
        ("Tour JSON/API", "Internal HTTP", "Scenes, hotspots, room and room-type availability, tour reservation submission.", "Validation, throttling, public-data minimization."),
        ("Cloudflare Tunnel", "Inbound transport", "Public HTTPS access to the local-hosted application.", "Trusted proxy/request-aware URL behavior."),
        ("Scheduler", "Internal", "Hourly alternative-offer expiration and reconciliation.", "Idempotent command execution and logs."),
    ], [1700, 1400, 3020, 3240])
    doc.add_heading("3.3 Communications", level=2)
    add_paragraph(doc, "Public production traffic shall use HTTPS. Local development may use HTTP on localhost/LAN. Webhook callbacks shall use the configured public HTTPS endpoint. Session cookies, CSRF controls, throttling, signed URLs, and anti-spam middleware shall be applied according to route risk.")
    doc.add_heading("3.4 Media and File Interfaces", level=2)
    add_paragraph(doc, "The system shall support uploaded branding, hero images, room images, panorama scenes, thumbnails, and hotspot media. Generated media URLs shall be compatible with both localhost and Cloudflare. Backup uploads shall be restricted to supported formats and authorized users.")

    doc.add_heading("4. Use-Case Catalog", level=1)
    doc.add_heading("4.1 Use Cases by Functional Area", level=2)
    add_table(doc, ["Functional area", "Use cases"], FUNCTIONAL_GROUPS, [3500, 5860])
    doc.add_heading("4.2 Use Cases by Stakeholder", level=2)
    add_table(doc, ["Stakeholder", "Primary use cases"], STAKEHOLDERS, [2800, 6560])

    doc.add_heading("5. Functional Requirements", level=1)
    add_paragraph(doc, "Each use case below is a functional requirement. Main scenarios describe the normal successful flow; validation, authorization, availability conflicts, expired links, gateway failures, and external-service failures shall produce clear errors without corrupting state.")
    for uc_id, title, area, actors, priority, frequency, summary, steps, notes in USE_CASES:
        doc.add_heading(f"{uc_id}: {title}", level=2)
        add_label_value(doc, "Functional area", area)
        add_label_value(doc, "Summary", summary)
        add_label_value(doc, "Priority", priority)
        add_label_value(doc, "Use frequency", frequency)
        add_label_value(doc, "Direct actors", actors)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        r = p.add_run("Main Success Scenario")
        set_font(r, size=10.5, bold=True, color=GREEN)
        for step in steps:
            add_number(doc, step)
        add_label_value(doc, "Notes and constraints", notes)

    doc.add_heading("6. Data Requirements", level=1)
    doc.add_heading("6.1 Core Data Entities", level=2)
    add_table(doc, ["Entity", "Purpose", "Key relationships / controls"], [
        ("User", "Staff/admin identity, role, and custom permissions.", "Links to reviews, room assignments, logs; password hashing and role enforcement."),
        ("Guest Account", "Guest identity, verification, contact, and account state.", "Owns reservations, feedback, support; may be disabled."),
        ("Reservation", "Central stay request and lifecycle record.", "Room requests, holds, assignments, guests, payments, charges, logs, offers, feedback."),
        ("Reservation Room Request", "Requested room type, quantity, and occupants.", "Supports mixed room types within one request."),
        ("Room Type / Room / Floor", "Sellable accommodation category and physical inventory.", "Amenities, rates, sharing/pricing model, holds and assignments."),
        ("Room Hold", "Temporary reservation of room inventory/capacity.", "Advance/offer lifecycle, expiry, release, conversion at check-in."),
        ("Room Assignment / Guest", "Actual checked-in room occupancy and persons.", "Capacity, timestamps, billing guest, remarks, check-in/out state."),
        ("Charge / Payment", "Financial obligations and receipts/gateway transactions.", "Totals, balance, official receipt, payment mode/status/reference."),
        ("Alternative Offer", "Time-limited alternative accommodation proposal.", "Offered rooms, expiry, guest response, reservation status."),
        ("Feedback", "Stay ratings, comments, review, consent and visibility.", "One per reservation; public eligibility enforced by model."),
        ("Support Inquiry / Reply", "Guest support case and conversation.", "Category, status, internal notes, handlers, guest/staff replies."),
        ("Tour Waypoint / Hotspot", "360 scene and interaction graph.", "Panorama/thumbnail/media, linked room type, action target and position."),
        ("Setting", "Configurable site, payment, discount and report values.", "Defaults plus stored overrides; permissions control modification."),
        ("Audit/Log/Snapshot", "Operational history and check-in evidence.", "Reservation logs, stay logs, snapshots, force-deletion evidence."),
    ], [1900, 3350, 4110])
    doc.add_heading("6.2 Data Validation", level=2)
    for text in ["Email addresses shall use valid format and uniqueness where required.", "Philippine mobile numbers shall accept supported 09, +639, or 639 formats and reject invalid lengths/patterns.", "Check-out shall be after check-in; past check-in requests shall be rejected under the guest date policy.", "Ratings shall be integers from 1 through 5; public room-type consent shall be ignored unless public-comment consent is also true.", "Amounts and percentages shall be numeric, nonnegative, and within configured bounds.", "Room allocation shall satisfy active status, date overlap, sharing model, room/bed capacity, assignments, and holds."]:
        add_bullet(doc, text)
    doc.add_heading("6.3 Retention, Privacy, and Audit", level=2)
    add_paragraph(doc, "Operational records shall be retained according to institutional policy and current application retention behavior. Tracking access for terminal reservations expires after defined periods. Public testimonials shall never expose names, email addresses, reservation references, room numbers, or stay dates. Privileged destructive actions and workflow transitions shall retain actor, timestamp, and reason where supported.")

    doc.add_heading("7. Business Rules", level=1)
    rules = [
        ("BR-01", "A stay form creates a pending request; it is not an immediate confirmed booking."),
        ("BR-02", "Approval requires revalidated inventory and creates advance holds before online payment becomes available."),
        ("BR-03", "Alternative offers are time-limited; acceptance converts the offered allocation, while decline/expiry releases it."),
        ("BR-04", "Only approved or confirmed reservations with active advance holds may accept guest online payment."),
        ("BR-05", "Successful online payment may confirm a reservation but must not bypass pending staff review."),
        ("BR-06", "Check-in converts validated holds/room selections into assignments and records guests, charges, payment, snapshot, and logs."),
        ("BR-07", "Feedback is available only to verified owners of checked-out reservations and only once per reservation."),
        ("BR-08", "Public testimonials require consent, written comment, reviewed status, review timestamp, and public visibility."),
        ("BR-09", "Room type appears in a testimonial only with separate room-type consent; all other identifying stay details remain private."),
        ("BR-10", "Super administrators inherit all permissions; custom permissions replace role defaults for other configured users."),
        ("BR-11", "Request-aware/relative URLs are preferred for internal links and media across localhost and Cloudflare."),
        ("BR-12", "Holds, assignments, and terminal status transitions must release or update inventory consistently."),
    ]
    add_table(doc, ["Rule", "Requirement"], rules, [1200, 8160])

    doc.add_heading("8. Non-Functional Requirements", level=1)
    nfrs = {
        "8.1 Usability": [
            "NFR-USE-01: The public and guest interfaces shall use clear request-based language and shall not imply immediate booking before staff review.",
            "NFR-USE-02: Forms shall provide labels, required-state cues, validation feedback, and preserved input after validation failure.",
            "NFR-USE-03: Reservation status pages shall provide plain-language next-step guidance and expected processing time where applicable.",
            "NFR-USE-04: Responsive layouts shall remain usable on common desktop and mobile viewport sizes.",
        ],
        "8.2 Performance": [
            "NFR-PERF-01: Public catalog and availability queries should return within 3 seconds under normal institutional load, excluding external network delay.",
            "NFR-PERF-02: Database queries shall use eager loading, indexes, bounded result sets, and aggregation appropriate to the requested report.",
            "NFR-PERF-03: Panorama and media assets should use thumbnails/lazy loading where applicable to reduce initial transfer cost.",
        ],
        "8.3 Reliability and Availability": [
            "NFR-REL-01: Reservation, hold, assignment, payment, and checkout workflows shall use transactional updates where partial state would be unsafe.",
            "NFR-REL-02: Scheduled expiration and webhook processing shall be idempotent and safe to retry.",
            "NFR-REL-03: External email or payment failures shall produce actionable logs and shall not corrupt the underlying reservation.",
            "NFR-REL-04: The system shall provide backup and restore facilities restricted to authorized administrators.",
        ],
        "8.4 Security and Privacy": [
            "NFR-SEC-01: Passwords shall be hashed; authenticated sessions shall use Laravel session, CSRF, and cookie protections.",
            "NFR-SEC-02: Role-based and custom permissions shall govern page access, fields, row actions, bulk actions, reports, and destructive operations.",
            "NFR-SEC-03: Public forms and sensitive lookup/payment endpoints shall use throttling and anti-spam controls appropriate to risk.",
            "NFR-SEC-04: Signed or time-limited links shall protect secure tracking, password reset, verification, alternative offers, and payment access where applicable.",
            "NFR-SEC-05: PayMongo webhook signatures shall be verified; secrets shall remain environment-managed.",
            "NFR-SEC-06: Public testimonials shall be anonymous and consent-driven, with separate consent for room-type disclosure.",
            "NFR-SEC-07: Guest account ownership shall be enforced before private reservation, feedback, and support data is displayed or modified.",
        ],
        "8.5 Maintainability": [
            "NFR-MNT-01: Reservation status labels, colors, and guidance shall use a centralized source of truth.",
            "NFR-MNT-02: Guest-site configuration shall use centralized defaults and stored overrides.",
            "NFR-MNT-03: Business workflows shall remain encapsulated in services for availability, holds, approval, check-in, alternatives, payments, and reporting.",
            "NFR-MNT-04: Database schema changes shall be versioned through reversible migrations.",
            "NFR-MNT-05: Automated tests shall cover core guest, reservation, payment, authorization, report, and privacy workflows.",
        ],
        "8.6 Compatibility and Portability": [
            "NFR-COMP-01: Localhost/LAN and Cloudflare-hosted operation shall be equally supported.",
            "NFR-COMP-02: Assets, media, internal links, redirects, callbacks, QR codes, and generated URLs shall not depend unnecessarily on one APP_URL value.",
            "NFR-COMP-03: The public interface shall support current standards-based browsers; advanced VR features may degrade gracefully when unsupported.",
        ],
        "8.7 Scalability": [
            "NFR-SCALE-01: New room types, rooms, floors, amenities, services, tour scenes, hotspots, users, and settings shall be addable without code changes.",
            "NFR-SCALE-02: Reporting and catalog queries shall remain filterable and bounded as records grow.",
            "NFR-SCALE-03: Queue, cache, session, and scheduler components should be configurable for larger deployments.",
        ],
        "8.8 Accessibility": [
            "NFR-A11Y-01: Interactive controls shall have programmatic labels and keyboard-operable semantics.",
            "NFR-A11Y-02: Color shall not be the sole status indicator; text labels and guidance shall accompany badges.",
            "NFR-A11Y-03: Images and key media controls shall provide meaningful alternative text or accessible names.",
            "NFR-A11Y-04: High-contrast and large-text guest settings shall remain functional where configured.",
        ],
        "8.9 Auditability and Observability": [
            "NFR-AUD-01: Status transitions, check-in/out, assignments, payments, review actions, and privileged destructive actions shall record actor and time where supported.",
            "NFR-AUD-02: Application, gateway, mail, scheduler, and backup errors shall be logged with enough context for diagnosis without exposing secrets.",
            "NFR-AUD-03: Reports shall derive from stored operational records and canonical status definitions.",
        ],
    }
    for heading, items in nfrs.items():
        doc.add_heading(heading, level=2)
        for item in items:
            add_bullet(doc, item)

    doc.add_heading("9. Traceability and Acceptance", level=1)
    doc.add_heading("9.1 Functional Traceability", level=2)
    add_table(doc, ["Requirement area", "Use cases", "Primary verification"], [
        (area, uc_range, "Feature/integration tests plus role-appropriate UI verification")
        for area, uc_range in FUNCTIONAL_GROUPS
    ], [3300, 2800, 3260])
    doc.add_heading("9.2 Acceptance Principles", level=2)
    for text in ["Every essential use case shall have at least one successful-path test and relevant validation/authorization tests.", "Permission-protected features shall be tested with allowed and denied roles or custom permissions.", "Availability, holds, assignments, alternative offers, and payments shall be tested for conflicts, expiry, and retry behavior.", "Public/guest features shall be checked on localhost and the Cloudflare-hosted target after deployment-sensitive changes.", "Public testimonial tests shall prove that nonconsenting, unreviewed, internal, or commentless feedback is excluded and identifying details are not rendered.", "Reports shall be reconciled against known fixture data and canonical status definitions."]:
        add_bullet(doc, text)

    doc.add_heading("Appendix A. Reservation Status Lifecycle", level=1)
    add_table(doc, ["Status", "Meaning", "Typical next states"], [
        ("Pending Review", "Submitted request awaiting staff decision.", "Alternative Offer Pending, Approved, Declined, Cancelled"),
        ("Alternative Offer Pending", "Alternative inventory is held while awaiting guest response.", "Confirmed/Approved as implemented, Pending Review, Declined/Expired"),
        ("Approved", "Staff approved the request and advance room holds exist.", "Confirmed, Cancelled, Checked In"),
        ("Confirmed", "Room space is reserved; payment or staff workflow confirmed the stay.", "Checked In, Cancelled"),
        ("Payment Pending", "Legacy/presentation status for an awaiting-payment state where applicable.", "Confirmed, Cancelled"),
        ("Declined", "Staff declined the request.", "Terminal; new request required"),
        ("Cancelled", "Reservation was cancelled by an authorized workflow.", "Terminal"),
        ("Checked In", "Guest and room assignments are active.", "Checked Out"),
        ("Checked Out", "Stay is complete and rooms are released.", "Terminal; feedback eligible for verified owner"),
    ], [1700, 4420, 3240])

    doc.add_heading("Appendix B. Glossary", level=1)
    add_table(doc, ["Term", "Definition"], [
        ("Advance hold", "A temporary reservation of room inventory created during approval before check-in."),
        ("Alternative offer", "A time-limited proposal for a different room type with associated held capacity."),
        ("Assignment", "The actual association of a checked-in guest with a physical room."),
        ("Guest account", "Authenticated self-service identity distinct from a staff/admin user."),
        ("Public testimonial", "Reviewed feedback explicitly approved for public display after guest consent."),
        ("Request-aware URL", "A URL generated from the current request/route context so it works across local and Cloudflare targets."),
        ("Room request line", "One requested room type, quantity, and occupant allocation within a reservation."),
        ("Shared room", "Accommodation sold by bed/capacity rather than exclusively as a private room."),
        ("Waypoint", "A panorama scene in the virtual-tour graph."),
        ("Hotspot", "An interactive point in a tour scene that navigates, displays media, or triggers another supported action."),
    ], [2100, 7260])

    # Keep headings with following content and avoid stranded labels.
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            p.paragraph_format.keep_with_next = True
        if p.text.startswith(("Summary:", "Priority:", "Use frequency:", "Direct actors:")):
            p.paragraph_format.keep_with_next = True

    doc.core_properties.title = "UHLMS Software Requirements Specification"
    doc.core_properties.subject = "Complete functional and non-functional requirements for the University Homestay Lodging Management System"
    doc.core_properties.author = "Central Mindanao University - UHLMS Project Team"
    doc.core_properties.keywords = "UHLMS, SRS, lodging management, reservations, virtual tour, requirements"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
